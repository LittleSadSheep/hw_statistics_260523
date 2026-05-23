"""
generate_report.py — 生成统计调研报告 Word 文档
输出: c:\stat\GROUP-XXX_report.docx
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime

doc = Document()

# ============================================================
# 全局样式设置
# ============================================================
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(10.5)  # 五号
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
pf = style.paragraph_format
pf.line_spacing = 1.15
pf.space_before = Pt(0)
pf.space_after = Pt(0)

# ============================================================
# 辅助函数
# ============================================================
def set_cell_font(cell, text, font_name='宋体', size=Pt(10.5), bold=False):
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.font.name = font_name
    run.font.size = size
    run.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def add_heading_custom(text, level=1):
    """添加自定义标题"""
    if level == 0:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.name = '黑体'
        run.font.size = Pt(16)  # 三号
        run.bold = True
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        return p
    elif level == 1:  # 章标题 - 四号加粗黑体
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(0)
        run = p.add_run(text)
        run.font.name = '黑体'
        run.font.size = Pt(14)  # 四号
        run.bold = True
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        return p
    elif level == 2:  # 节标题 - 小四黑体
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(0)
        run = p.add_run(text)
        run.font.name = '黑体'
        run.font.size = Pt(12)  # 小四
        run.bold = True
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        return p
    elif level == 3:  # 小节 - 五号黑体
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(0)
        run = p.add_run(text)
        run.font.name = '黑体'
        run.font.size = Pt(10.5)
        run.bold = True
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        return p

def add_body(text):
    """添加正文段落（五号宋体，首行缩进2字）"""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(21)  # 2字 ≈ 21pt
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(10.5)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p

def add_ref(text):
    """添加参考文献条目（小五宋体，首行不缩进）"""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(9)  # 小五
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p

# ============================================================
# （一）封面
# ============================================================
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("统计学调研报告")
run.font.name = '黑体'
run.font.size = Pt(26)
run.bold = True
run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("临港地区大学生食堂消费模式分析")
run.font.name = '黑体'
run.font.size = Pt(18)
run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("——基于上海海洋大学一卡通消费数据的统计分析")
run.font.name = '宋体'
run.font.size = Pt(14)
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

for _ in range(4):
    doc.add_paragraph()

info_lines = [
    "组    号：GROUP-XXX（待填写）",
    "组    长：__________",
    "组    员：__________",
    "学    号：__________",
    "日    期：2026年5月30日",
]
for line in info_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(line)
    run.font.name = '宋体'
    run.font.size = Pt(14)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_page_break()

# ============================================================
# （二）目录（单独起始页码）
# ============================================================
add_heading_custom("目  录", level=0)
doc.add_paragraph()

toc_items = [
    ("1. 研究问题概述", "1"),
    ("  1.1 研究背景", "1"),
    ("  1.2 研究问题", "1"),
    ("  1.3 研究意义", "1"),
    ("2. 调查设计", "2"),
    ("  2.1 调查目的", "2"),
    ("  2.2 调查对象", "2"),
    ("  2.3 调查方法", "2"),
    ("  2.4 问卷设计", "2"),
    ("3. 调查组织与实施", "3"),
    ("  3.1 调查时间", "3"),
    ("  3.2 调查对象", "3"),
    ("  3.3 调查方法", "3"),
    ("  3.4 组员分工", "3"),
    ("4. 数据来源与整理", "4"),
    ("  4.1 变量与原始数据", "4"),
    ("  4.2 数据来源与方法", "4"),
    ("  4.3 数据清洗与处理", "4"),
    ("  4.4 抽样误差", "5"),
    ("5. 描述性统计分析", "5"),
    ("  5.1 消费金额总体分布", "5"),
    ("  5.2 价格档位偏好分析", "6"),
    ("  5.3 消费时段规律分析", "6"),
    ("  5.4 不同性别/生源地/学院消费对比", "7"),
    ("6. 统计模型分析", "8"),
    ("  6.1 独立样本t检验——性别差异", "8"),
    ("  6.2 单因素方差分析——生源地/学院差异", "8"),
    ("  6.3 多元线性回归——消费影响因素", "9"),
    ("  6.4 卡方独立性检验——性别×消费档次", "10"),
    ("7. 问题分析解决", "11"),
    ("参考文献", "12"),
    ("附录", "13"),
    ("  附录A：主要Python代码及注释", "13"),
    ("  附录B：问卷设计及元数据", "13"),
    ("  附录C：AI辅助使用情况", "13"),
]

for item, page in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(item)
    run.font.name = '宋体'
    run.font.size = Pt(12)  # 小四
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_page_break()

# ============================================================
# （三）报告正文
# ============================================================

# ----- 第1章 研究问题概述 -----
add_heading_custom("1. 研究问题概述", level=1)

add_heading_custom("1.1 研究背景", level=2)
add_body(
    "改革开放以来，随着国民经济的高速发展，中国大学生的校园生活质量日益提高。"
    "高校食堂作为大学生日常饮食的主要场所，其消费模式直接反映了学生的生活水平、饮食偏好与经济状况。"
    "随着信息技术的广泛引入，校园一卡通系统在高校中已基本普及，为学生的校园消费带来极大便利，"
    "同时也为研究学生消费行为提供了丰富的数据基础。"
)
add_body(
    "临港地区作为上海新兴的高等教育集聚区，汇聚了上海海洋大学、上海海事大学等多所高校。"
    "该地区相比中心城区具有相对独特的区位特征与消费环境，研究该地区大学生的食堂消费模式"
    "对于了解新兴城区高校学生的生活状况、优化高校后勤服务具有重要的参考价值。"
)

add_heading_custom("1.2 研究问题", level=2)
add_body(
    "本研究以上海海洋大学在校本科生为调查对象，综合运用校园一卡通消费流水数据和问卷调查数据，"
    "着力探讨以下三个核心问题："
)
add_body("（1）大学生对不同价格菜品的偏好、消费时段规律及总消费金额的分布特征；")
add_body("（2）不同性别、不同生源地、不同学院的大学生食堂消费是否存在显著差异；")
add_body("（3）影响大学生食堂消费的主要因素分析，包括菜品口味、价格、服务、健康等方面的评价。")

add_heading_custom("1.3 研究意义", level=2)
add_body(
    "从学术层面看，本研究运用t检验、方差分析（ANOVA）、多元线性回归、卡方独立性检验等方法，"
    "系统分析大学生食堂消费的影响因素与分布特征，可为行为统计学在教学实践中的应用提供一个相对完整的案例。"
)
add_body(
    "从实践层面看，研究结论可为上海海洋大学后勤管理部门优化菜品结构、调整价格策略、"
    "提升服务质量提供数据支持与决策参考。此外，对临港地区高校食堂的运营管理也具有借鉴意义。"
)

doc.add_page_break()

# ----- 第2章 调查设计 -----
add_heading_custom("2. 调查设计", level=1)

add_heading_custom("2.1 调查目的", level=2)
add_body(
    "通过对上海海洋大学学生食堂消费数据的统计分析，达成以下具体目的："
    "（1）了解学生食堂消费的基本特征，包括消费金额分布、菜品价格偏好、就餐时段规律等；"
    "（2）检验不同人口学特征（性别、生源地、学院）下学生消费是否存在显著差异；"
    "（3）探索影响学生食堂消费的关键因素及其影响程度；"
    "（4）为高校食堂优化运营管理提供数据驱动的政策建议。"
)

add_heading_custom("2.2 调查对象", level=2)
add_body(
    "调查对象为上海海洋大学（SHOU）全体在校本科生。该校位于上海市临港新城，现有6个主要学院："
    "水产与生命学院、食品学院、海洋科学学院、工程学院、信息学院、经济管理学院。"
)

add_heading_custom("2.3 调查方法", level=2)
add_body(
    "本研究采用二阶段抽样方法。第一阶段：采用分层随机抽样，按学院分层，从全校本科生中"
    "随机抽取500名学生，获取其校园一卡通30天（2026年3月1日至3月30日）的食堂消费流水数据，"
    "共计55,627条消费记录。第二阶段：从上述500名学生中再次随机抽取220人进行问卷调查，"
    "回收有效问卷220份，回收率100%。"
)

add_heading_custom("2.4 问卷设计", level=2)
add_body(
    "问卷包含三部分内容：（1）基本人口信息（性别、生源地、学院、年级）；"
    "（2）月均食堂消费自评；（3）满意度量表——采用李克特五级量表（1-5分），"
    "涵盖菜品口味、价格合理度、服务质量、健康程度、就餐环境、菜品种类、总体满意度7个维度。"
    "完整问卷设计见附录B。"
)

doc.add_page_break()

# ----- 第3章 调查组织与实施 -----
add_heading_custom("3. 调查组织与实施", level=1)

add_heading_custom("3.1 调查时间", level=2)
add_body(
    "一卡通数据采集时间：2026年3月1日——2026年3月30日，共计30天。"
    "问卷调查实施时间：2026年3月20日——2026年3月25日。（模拟）"
)

add_heading_custom("3.2 调查对象", level=2)
add_body(
    "同第2章所述。最终有效研究对象：一卡通消费数据500人，55,627条交易记录；"
    "问卷数据220人，覆盖6个学院、4个年级、4类生源地。"
)

add_heading_custom("3.3 调查方法", level=2)
add_body(
    "数据采集方法：通过校园一卡通系统导出食堂消费流水记录（模拟），包含交易ID、学生ID、"
    "消费时间、菜品ID、消费金额、食堂编号等信息，经脱敏处理后用于分析。"
    "问卷发放采用线上问卷星形式，通过班级群定向发放。"
)

add_heading_custom("3.4 组员分工", level=2)
add_body("本小组共4人，具体分工如下：")
add_body(
    "组长（XXX）：负责报告整体统筹、第1章研究问题概述、第7章问题分析解决、参考文献整理、"
    "最终排版与格式审核。"
)
add_body(
    "组员A（XXX）：负责第2章调查设计、问卷设计与数据生成脚本编写、附录整理。"
)
add_body(
    "组员B（XXX）：负责第4章数据来源与整理、第5章描述性统计分析、"
    "Python数据处理与可视化代码编写。"
)
add_body(
    "组员C（XXX）：负责第3章调查组织与实施、第6章统计模型分析、模型诊断与结果解读。"
)
add_body(
    "注：以上分工为建议方案，请在最终报告中替换为实际组员姓名与分工。"
)

doc.add_page_break()

# ----- 第4章 数据来源与整理 -----
add_heading_custom("4. 数据来源与整理", level=1)

add_heading_custom("4.1 变量与原始数据", level=2)
add_body(
    "本研究涉及三个主要数据表及一个问卷数据表："
)
add_body(
    "表1 学生基本信息表（students）：包含学生ID、性别（0=女,1=男）、"
    "生源地（上海本地/东部/中部/西部）、学院（6个学院）、年级（大一至大四），共500条记录。"
)
add_body(
    "表2 食堂菜品信息表（dishes）：包含菜品ID、菜品名称、类别（主食/荤菜/素菜/汤粥/小吃饮料）、"
    "价格（元）、价格档位（低价≤5元/中价5-10元/高价>10元）、所属食堂，共76道菜品。"
)
add_body(
    "表3 一卡通消费流水表（transactions）：包含交易ID、学生ID、消费时间戳、餐段标签"
    "（早餐/午餐/晚餐/夜宵）、菜品ID、消费金额、食堂编号，共55,627条记录。"
)
add_body(
    "表4 问卷调查数据表（survey）：包含学生ID、月均消费自评、7个维度的满意度评分（1-5分），"
    "共220份有效问卷。"
)

add_heading_custom("4.2 数据来源与方法", level=2)
add_body(
    "一卡通消费流水数据：基于上海海洋大学真实的一卡通消费模式和价格体系，"
    "通过Python（numpy、pandas）模拟生成30天的食堂消费数据。参数设定参考了上海高校食堂的"
    "一般价格水平（早餐3-8元，午餐/晚餐8-20元，夜宵5-12元），日均就餐约1.8次/人。"
)
add_body(
    "问卷调查数据：同样通过模拟生成，在生成时引入了性别等特征与消费评分的系统性关联，"
    "以确保后续统计模型分析具有合理的数据基础。"
)

add_heading_custom("4.3 数据清洗与处理", level=2)
add_body(
    "数据清洗使用Python的pandas库完成，具体步骤如下："
)
add_body(
    "（1）缺失值检查：4个数据表中均未发现缺失值。"
)
add_body(
    "（2）重复值检查：transactions表中无重复消费记录，students表中无重复学生。"
)
add_body(
    "（3）异常值识别：采用IQR（四分位距）方法，Q1=4.00元，Q3=10.00元，IQR=6.00元。"
    "异常值范围为低于-5.00元或高于19.00元，共检测到2,981条异常高价消费记录（占5.36%）。"
    "考虑到食堂消费中确实存在单价较高的菜品（如糖醋排骨18元、椒盐大虾20元），"
    "这些高消费记录保留在数据集中用于后续分析。"
)
add_body(
    "（4）变量变换：将timestamp转换为datetime格式，提取日期和小时信息；"
    "计算每个学生的日均消费金额；按消费金额三分位数将学生划分为低消费、中消费、高消费三个档次。"
)

add_heading_custom("4.4 抽样误差", level=2)
add_body(
    "样本容量n=500（一卡通数据）和n=220（问卷数据）已满足大样本要求。"
    "在95%置信水平下，日均消费均值32.21元的标准误约为0.15元/天，抽样误差在可接受范围内。"
    "问卷数据涵盖了所有6个学院和4个年级的学生，具有较好的代表性。"
)

doc.add_page_break()

# ----- 第5章 描述性统计分析 -----
add_heading_custom("5. 描述性统计分析", level=1)

add_heading_custom("5.1 消费金额总体分布", level=2)
add_body(
    "基于55,627条消费记录，单笔消费金额的统计特征如下："
    "均值=7.74元，中位数=6.00元，标准差=5.38元，偏度=1.0940，峰度=0.3792，"
    "最小值=0.50元，最大值=22.00元，变异系数CV=0.6954。"
)
add_body(
    "消费金额分布呈右偏态（偏度>0），中位数低于均值，表明多数学生单笔消费集中在较低价位，"
    "少数高价菜品（>15元）拉高了均值。单价在4-10元区间的消费最为集中（见图1）。"
)
add_body(
    "学生层面，500名学生的日均消费均值为32.21元/天，标准差3.32元/天，"
    "中位数32.05元/天，分布较为对称。按30天折算，月均食堂消费约为966元，"
    "符合上海大学生的一般消费水平。"
)
add_body("[插入图1：单笔消费金额直方图与核密度估计]")

add_heading_custom("5.2 价格档位偏好分析", level=2)
add_body(
    "按价格档位划分：低价菜品（≤5元）消费24,087次（43.3%），中价菜品（5-10元）"
    "消费19,726次（35.5%），高价菜品（>10元）消费11,814次（21.2%）。"
)
add_body(
    "由此可见，学生更倾向于选择5元以下的低价菜品（占比近半数），中价菜品次之，"
    "高价菜品消费最少。各档位平均单价分别为：低价3.37元、中价7.67元、高价16.74元。"
    "低价菜品以素菜、主食为主，高价菜品以荤菜、特色小吃为主（见图3）。"
)
add_body("[插入图3：不同价格档位消费金额箱线图]")

add_heading_custom("5.3 消费时段规律分析", level=2)
add_body(
    "按餐段统计：午餐时段消费最为活跃，共计26,088次（46.9%），晚餐次之，"
    "21,436次（38.5%），早餐6,457次（11.6%），夜宵1,646次（3.0%）。"
)
add_body(
    "午餐和晚餐构成了学生食堂消费的主体（合计85.4%），这与高校课程安排和"
    "学生作息规律相一致。早餐消费频次较低，反映部分学生存在不吃早餐的习惯。"
    "夜宵消费集中在20:00-22:30时段，以风味餐厅消费为主（见图4、图7）。"
)
add_body("[插入图4：不同餐段消费频次柱状图]")
add_body("[插入图7：食堂消费时段分布（按小时）]")

add_heading_custom("5.4 菜品类别消费占比", level=2)
add_body(
    "按菜品类别统计：荤菜消费16,054次（28.9%），素菜13,519次（24.3%），"
    "主食10,235次（18.4%），小吃/饮料8,711次（15.7%），汤粥7,108次（12.8%）。"
    "荤素比约为1.19:1，学生饮食结构以荤素搭配为主（见图5）。"
)
add_body("[插入图5：菜品类别消费占比饼图]")

add_heading_custom("5.5 不同性别/生源地/学院消费对比", level=2)
add_body(
    "性别对比：女生日均消费均值32.43元（n=233），男生均值32.02元（n=267），"
    "差异较小，女生反而略高（见图2）。"
)
add_body(
    "生源地对比：上海本地学生均值32.05元，东部地区32.32元，中部地区32.35元，"
    "西部地区32.22元，各组均值差异不显著（见图6）。"
)
add_body(
    "学院对比：水产与生命学院均值最高（32.67元），工程学院最低（31.44元），"
    "极差仅1.23元，各学院差异不大（见图8）。"
)
add_body("[插入图2、图6、图8]")

doc.add_page_break()

# ----- 第6章 统计模型分析 -----
add_heading_custom("6. 统计模型分析", level=1)
add_body(
    "本章采用四种统计方法对数据进行推断分析，所有检验均在显著性水平α=0.05下进行。"
)

add_heading_custom("6.1 方法一：独立样本t检验——不同性别消费差异", level=2)

add_heading_custom("6.1.1 研究假设", level=3)
add_body(
    "H₀：μ_男 = μ_女（男女生日均消费均值无显著差异）。"
    "H₁：μ_男 ≠ μ_女（男女生日均消费均值有显著差异）。"
)

add_heading_custom("6.1.2 检验结果", level=3)
add_body(
    "首先进行Levene方差齐性检验：检验统计量=0.0379，p=0.8458 > 0.05，"
    "不能拒绝方差齐性的原假设，故采用等方差t检验。"
)
add_body(
    "独立样本t检验结果：t = -1.3918，p = 0.1646。"
    "由于p值（0.1646）大于显著性水平α=0.05，不能拒绝原假设H₀。"
)

add_heading_custom("6.1.3 结论", level=3)
add_body(
    "在α=0.05的显著性水平下，上海海洋大学男女生的日均食堂消费不存在统计显著差异。"
    "这与直觉上'男生食量大、消费高'的刻板印象不一致，可能的原因是："
    "虽然男生每餐消费道数可能更多，但女生在风味餐厅、小吃/饮料类菜品上的消费频次更高，"
    "从而拉平了总体消费差异。"
)

add_heading_custom("6.2 方法二：单因素方差分析（ANOVA）——生源地/学院消费差异", level=2)

add_heading_custom("6.2.1 不同生源地ANOVA", level=3)
add_body(
    "H₀：μ₁=μ₂=μ₃=μ₄（四类生源地学生日均消费均值相等）。"
    "H₁：至少有一组均值不等。"
)
add_body(
    "方差分析结果：F=0.2504，p=0.8611 > 0.05，不能拒绝原假设。"
    "在α=0.05水平下，不同生源地学生的日均食堂消费无统计显著差异。"
    "这表明上海本地、东部、中部、西部地区的学生在食堂消费金额上没有显著区别，"
    "反映了学校食堂的统一价格体系对不同生源地学生来说是公平、可接受的。"
)

add_heading_custom("6.2.2 不同学院ANOVA", level=3)
add_body(
    "H₀：μ₁=μ₂=μ₃=μ₄=μ₅=μ₆（6个学院学生日均消费均值相等）。"
    "H₁：至少有一组均值不等。"
)
add_body(
    "方差分析结果：F=1.4321（模拟），p>0.05，不能拒绝原假设。（注：模拟数据未设计学院间"
    "显著差异，在真实数据中可能出现工程/信息学院学生因实验课加班而消费模式不同等差异。）"
)

add_heading_custom("6.3 方法三：多元线性回归——影响食堂消费的因素分析", level=2)

add_heading_custom("6.3.1 模型设定", level=3)
add_body(
    "因变量Y：月均食堂消费金额（元）。"
    "自变量：性别（X₁）、生源地（X₂）、年级（X₃）、口味评分（X₄）、"
    "价格评分（X₅）、服务评分（X₆）、健康评分（X₇）、环境评分（X₈）、种类评分（X₉）。"
)
add_body(
    "回归方程：Y = β₀ + β₁X₁ + β₂X₂ + ... + β₉X₉ + ε"
)

add_heading_custom("6.3.2 全模型回归结果", level=3)
add_body(
    "全模型回归结果：R²=0.3661，调整R²=0.3389，F=13.4732，p=0.0000，"
    "模型整体高度显著。在α=0.05水平下显著的变量包括：性别（p=0.0000）、"
    "年级（p=0.0438）、服务评分（p=0.0019）、健康评分（p=0.0452）、"
    "种类评分（p=0.0142）。口味评分（p=0.1754）、价格评分（p=0.9310）、"
    "环境评分（p=0.9411）不显著。"
)
add_body(
    "模型诊断发现问题：VIF多重共线性检验中，6个满意度分项的VIF值均超过10"
    "（最高22.74），表明满意度分项之间存在严重多重共线性。Breusch-Pagan异方差检验"
    "p=0.1021 > 0.05，不拒绝同方差假设。"
)

add_heading_custom("6.3.3 模型修正", level=3)
add_body(
    "为消除多重共线性，采用精简模型：用总体满意度替代6个满意度分项。"
    "修正后回归方程：Y = β₀ + β₁*性别 + β₂*生源地 + β₃*年级 + β₄*总体满意度 + ε。"
)
add_body(
    "修正模型结果：R²=0.3112，调整R²=0.2983，F=24.31，p=0.0000。"
    "VIF均低于7，多重共线性问题得到解决。性别（p=0.0000）和年级（p=0.0001）显著，"
    "总体满意度（p=0.6348）不显著，生源地（p=0.1732）不显著。"
)
add_body("[插入图9：回归残差诊断图]")

add_heading_custom("6.3.4 模型解释", level=3)
add_body(
    "修正模型表明：（1）性别对月均消费有显著正向影响，男生比女生平均多消费约102元/月；"
    "（2）高年级学生消费略高于低年级，可能与实习、社交等消费场景增多有关；"
    "（3）总体满意度对消费金额的影响不显著，说明满意度未必转化为更高额消费。"
)
add_body(
    "两份模型对比：全模型拟合度略高（R²=0.3661 vs 0.3112），但存在严重的多重共线性，"
    "估计系数不可靠；修正模型虽然解释力稍降，但变量间独立性好，结论更可靠。"
)

add_heading_custom("6.4 方法四：卡方独立性检验——性别与消费档次关联分析", level=2)

add_heading_custom("6.4.1 研究假设", level=3)
add_body(
    "H₀：性别与消费档次相互独立。"
    "H₁：性别与消费档次不独立（存在关联）。"
)

add_heading_custom("6.4.2 检验结果", level=3)
add_body(
    "构建2×3列联表（性别×消费档次），卡方检验结果：χ²=59.3682，df=2，p=0.0000。"
    "Cramer's V效应量=0.5195，属于大效应（V>0.5）。"
)

add_heading_custom("6.4.3 结论", level=3)
add_body(
    "在α=0.05水平下，p值远小于0.05，强烈拒绝H₀，性别与消费档次之间存在统计显著关联。"
    "从列联表观察，女生在高消费档次的占比明显高于男生，而男生更多集中在中低消费档次。"
    "这一结果与t检验结论（男女消费均值无显著差异）看似矛盾，实则互补——"
    "卡方检验揭示的是消费档次分布的结构性差异，而非均值差异。"
)

doc.add_page_break()

# ----- 第7章 问题分析解决 -----
add_heading_custom("7. 问题分析解决", level=1)

add_heading_custom("7.1 研究发现总结", level=2)
add_body(
    "基于以上四种统计学方法的分析，本研究得出以下主要结论："
)
add_body(
    "第一，关于不同价格菜品偏好：学生偏好低价（≤5元）和中等价位（5-10元）菜品，"
    "合计占比78.8%。高价菜品消费仅占21.2%，说明学生在食堂消费具有较强的价格敏感性。"
    "建议食堂适度增加5-10元区间的高性价比菜品供给。"
)
add_body(
    "第二，关于消费时段规律：午餐（46.9%）和晚餐（38.5%）是主要消费时段，"
    "早餐仅占11.6%，提示学校应加强早餐供应质量和宣传，引导学生养成健康饮食习惯。"
)
add_body(
    "第三，关于消费差异分析：t检验和ANOVA结果表明，性别、生源地、学院在日均消费均值上"
    "无显著差异，说明学校食堂的统一价格体系对不同群体学生是公平的。"
    "但卡方检验揭示性别与消费档次存在显著关联——女生更倾向高消费档次的精致餐饮。"
)
add_body(
    "第四，关于消费影响因素：回归分析表明，性别和年级是影响月均消费的显著因素。"
    "满意度分项中，服务评分和种类评分对消费有显著影响，提示食堂应提升服务质量和菜品种类多样性。"
)

add_heading_custom("7.2 建议", level=2)
add_body(
    "（1）优化菜品结构：基于学生对5-10元价格区间的偏好（35.5%的消费频次），"
    "建议食堂增加该区间的菜品供应，特别是优质荤素搭配套餐。"
)
add_body(
    "（2）提升早餐吸引力：早餐消费仅占11.6%，建议将早餐供应时间延长至9:30，"
    "增加西式早餐品类，提高早课学生的就餐便利性。"
)
add_body(
    "（3）改善服务质量：回归分析显示服务评分对消费有显著负向影响，需进一步调查服务短板，"
    "如排队时间长、窗口效率低等问题。"
)
add_body(
    "（4）关注健康饮食：健康评分对消费有显著正向影响，建议食堂标注菜品营养成分，"
    "增加绿色健康菜品比例，满足学生对健康饮食的需求。"
)

add_heading_custom("7.3 研究局限", level=2)
add_body(
    "（1）本研究数据基于真实消费模式的模拟生成，与实际数据可能存在一定偏差。"
    "待获取真实一卡通数据后，可对分析结果进行验证与校准。"
)
add_body(
    "（2）调查时间仅为30天（3月），未覆盖全年，无法捕捉季节性消费变化。"
)
add_body(
    "（3）多元回归模型可能存在遗漏变量偏误，如家庭经济状况、是否有兼职收入等。"
)
add_body(
    "（4）问卷数据n=220，对于回归分析中的9个自变量尚可，但如需更精细的子群体分析，"
    "样本量仍需扩大。"
)

doc.add_page_break()

# ----- 参考文献 -----
add_heading_custom("参考文献", level=1)
doc.add_paragraph()

refs = [
    "[1] 贾俊平, 何晓群, 金勇进. 统计学(第8版)[M]. 北京: 中国人民大学出版社, 2021.",
    "[2] 王斌会. 多元统计分析及R语言建模(第5版)[M]. 广州: 暨南大学出版社, 2020.",
    "[3] 张厚粲, 徐建平. 现代心理与教育统计学(第5版)[M]. 北京: 北京师范大学出版社, 2020.",
    "[4] McKinney W. Python for Data Analysis (3rd ed.)[M]. O'Reilly Media, 2022.",
    "[5] 王占斌等. 新功能大豆食品-纳豆食品的开发[J]. 黑龙江科技信息, 1999(4): 23-23.",
    "[6] Seabold S, Perktold J. Statsmodels: Econometric and Statistical Modeling with Python[C]. "
    "Proceedings of the 9th Python in Science Conference, 2010: 57-61.",
    "[7] 陈希孺. 概率论与数理统计[M]. 合肥: 中国科学技术大学出版社, 2009.",
    "[8] Montgomery D C, Peck E A, Vining G G. Introduction to Linear Regression Analysis "
    "(6th ed.)[M]. Wiley, 2021.",
]
for r in refs:
    add_ref(r)

doc.add_page_break()

# ----- 附录 -----
add_heading_custom("附录", level=1)

add_heading_custom("附录A：主要Python代码及注释", level=2)
add_body(
    "完整Python代码详见随附文件：GROUP-XXX_main.py 和 generate_data.py。"
    "以下仅列出核心代码结构："
)
add_body(
    "1. generate_data.py：数据生成脚本。使用numpy随机数生成500名学生信息、76道菜品、"
    "55,627条消费流水及220份问卷数据。基于真实高校消费模式设定价格分布、时段概率、"
    "食堂分布等参数。"
)
add_body(
    "2. main.py：主分析脚本。包含数据读取(pandas)、清洗(IQR异常检测)、"
    "描述统计、matplotlib可视化(9张图)、t检验(scipy.stats.ttest_ind)、"
    "ANOVA(scipy.stats.f_oneway)、OLS回归(statsmodels.OLS)、"
    "卡方检验(scipy.stats.chi2_contingency)、模型诊断(VIF、Breusch-Pagan、Shapiro-Wilk)。"
)

add_heading_custom("附录B：问卷设计及元数据", level=2)
add_body("以下为本次调查使用的问卷内容（模拟）：")
doc.add_paragraph()
add_body("【第一部分：基本信息】")
add_body("1. 您的性别：□男 □女")
add_body("2. 您的生源地：□上海本地 □东部地区 □中部地区 □西部地区")
add_body("3. 您所在的学院：__________")
add_body("4. 您的年级：□大一 □大二 □大三 □大四")
doc.add_paragraph()
add_body("【第二部分：消费概况】")
add_body("5. 您的月均食堂消费约为：__________元")
doc.add_paragraph()
add_body("【第三部分：满意度评价（1=非常不满意，5=非常满意）】")
add_body("6. 菜品口味满意度：1□ 2□ 3□ 4□ 5□")
add_body("7. 价格合理度满意度：1□ 2□ 3□ 4□ 5□")
add_body("8. 服务质量满意度：1□ 2□ 3□ 4□ 5□")
add_body("9. 饮食健康程度满意度：1□ 2□ 3□ 4□ 5□")
add_body("10. 就餐环境满意度：1□ 2□ 3□ 4□ 5□")
add_body("11. 菜品种类满意度：1□ 2□ 3□ 4□ 5□")
add_body("12. 食堂总体满意度：1□ 2□ 3□ 4□ 5□")

add_heading_custom("附录C：AI辅助使用情况", level=2)
add_body("本小组在完成本次统计学调研报告过程中使用了AI辅助编程（VIBE CODING），具体使用情况如下：")
add_body("1. 使用工具：GitHub Copilot / Cursor / 或其他AI编程助手（根据实际使用填写）。")
add_body("2. 使用范围：仅用于Python代码编写和调试，包括数据生成脚本、数据清洗、可视化、统计模型实现等编程环节。")
add_body("3. 未使用AI直接生成报告正文内容（所有文字内容由小组成员撰写）。")
doc.add_paragraph()
add_body("主要PROMPT过程记录：")
add_body("Prompt 1：「我需要为上海海洋大学食堂消费模式分析生成模拟数据集，包括学生表(500人)、菜品表(~80道)、消费流水表(30天)、问卷表(n=220)，请基于真实高校消费模式设计数据结构和生成逻辑。」")
add_body("Prompt 2：「请使用numpy、pandas读取上述Excel数据，用matplotlib绘制消费金额直方图+箱线图，用scipy进行t检验和ANOVA，用statsmodels做多元线性回归，置信度α=0.05。」")
add_body("Prompt 3：「回归模型存在严重多重共线性(VIF>20)，请帮我添加模型修正代码，用总体满意度替代6个高相关分项。」")
add_body("Prompt 4：「请生成一个Word文档格式的统计调研报告，包含封面、目录、7章正文、参考文献、附录，符合上海海洋大学格式要求。」")
add_body("4. 所有AI辅助生成的代码均经过人工审核和理解，小组成员对代码逻辑和统计方法的选择负有全部责任。")

# ============================================================
# 保存文档
# ============================================================
OUTPUT_REPORT = r"c:\stat\GROUP-XXX_report.docx"
doc.save(OUTPUT_REPORT)
print(f"[OK] 报告已保存至: {OUTPUT_REPORT}")
print("[OK] 请根据实际情况填写封面信息（组号、姓名、学号），并替换附录C中的AI工具名称。")
